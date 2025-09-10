#!/usr/bin/env python3
"""
Word Document Processor using python-docx

This module implements specialized processing for Word documents containing
banking data, with table extraction and text analysis capabilities.
"""

import logging
import time
import os
from typing import List, Dict, Optional, Tuple, Any, Union
from dataclasses import dataclass
import re
from pathlib import Path
from datetime import datetime

# python-docx imports
try:
    from docx import Document
    from docx.table import Table as DocxTable
    from docx.text.paragraph import Paragraph
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    # Create dummy classes for type hints when not available
    class Document:
        pass

# Alternative docx2txt for fallback
try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
except ImportError:
    DOCX2TXT_AVAILABLE = False


@dataclass
class WordTable:
    """Structure for Word document tables"""
    data: List[List[str]]
    row_count: int
    col_count: int
    confidence: float
    table_index: int
    has_header: bool = False


@dataclass
class WordProcessingResult:
    """Result structure for Word document processing"""
    success: bool
    text_content: str
    tables: List[WordTable]
    transactions: List[Dict[str, Any]]
    processing_time: float
    metadata: Dict[str, Any]
    error_message: Optional[str] = None


class WordProcessor:
    """
    Specialized processor for Word documents containing banking data.
    Uses python-docx for direct extraction of text and tables without conversion.
    """
    
    def __init__(self, debug: bool = False):
        """
        Initialize the Word Processor.
        
        Args:
            debug: Enable debug logging
        """
        self.debug = debug
        self.logger = self._setup_logger()
        
        # Banking patterns for text analysis
        self.banking_patterns = {
            'date': [
                r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b',
                r'\b\d{2,4}[/-]\d{1,2}[/-]\d{1,2}\b',
                r'\b\d{1,2}\s+(?:enero|febrero|marzo|abril|mayo|junio|julio|agosto|septiembre|octubre|noviembre|diciembre)\s+\d{2,4}\b',
                r'\b(?:january|february|march|april|may|june|july|august|september|october|november|december)\s+\d{1,2},?\s+\d{2,4}\b'
            ],
            'amount': [
                r'[-+]?\$?\s*\d{1,3}(?:[.,]\d{3})*(?:[.,]\d{2})?',
                r'[-+]?\d{1,3}(?:[.,]\d{3})*(?:[.,]\d{2})?\s*€',
                r'[-+]?\d+(?:[.,]\d{2})?'
            ],
            'transaction_keywords': [
                r'\b(?:transferencia|transfer|pago|payment|cargo|charge|abono|deposit|retiro|withdrawal)\b',
                r'\b(?:débito|debit|crédito|credit|saldo|balance|movimiento|movement)\b',
                r'\b(?:comisión|commission|fee|interés|interest|cajero|atm)\b'
            ]
        }
        
        # Table detection patterns
        self.table_headers = [
            # Spanish headers
            r'fecha', r'descripci[oó]n', r'importe', r'saldo', r'referencia',
            r'concepto', r'movimiento', r'debe', r'haber', r'tipo',
            # English headers
            r'date', r'description', r'amount', r'balance', r'reference',
            r'concept', r'movement', r'debit', r'credit', r'type'
        ]
        
        # Quality thresholds
        self.quality_thresholds = {
            'min_table_rows': 3,
            'min_table_cols': 2,
            'min_banking_keywords': 2,
            'min_date_matches': 1,
            'min_amount_matches': 1
        }
        
        if not PYTHON_DOCX_AVAILABLE and not DOCX2TXT_AVAILABLE:
            self.logger.warning("Neither python-docx nor docx2txt available. Install with: pip install python-docx docx2txt")
        
        self.logger.info("WordProcessor initialized")
    
    def _setup_logger(self) -> logging.Logger:
        """Set up logger with consistent formatting"""
        logger = logging.getLogger(f"{__name__}.WordProcessor")
        
        if not logger.handlers:
            handler = logging.StreamHandler()
            formatter = logging.Formatter(
                '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
            )
            handler.setFormatter(formatter)
            logger.addHandler(handler)
        
        logger.setLevel(logging.DEBUG if self.debug else logging.INFO)
        return logger
    
    def process_word(self, file_path: str) -> WordProcessingResult:
        """
        Process Word document and extract banking data.
        
        Args:
            file_path: Path to Word document (.docx)
            
        Returns:
            WordProcessingResult with extracted content and transactions
        """
        if not Path(file_path).exists():
            return WordProcessingResult(
                success=False,
                text_content="",
                tables=[],
                transactions=[],
                processing_time=0.0,
                metadata={},
                error_message=f"File not found: {file_path}"
            )
        
        start_time = time.time()
        
        try:
            self.logger.info(f"Processing Word document: {file_path}")
            
            # Try python-docx first (preferred method)
            if PYTHON_DOCX_AVAILABLE:
                result = self._process_with_python_docx(file_path)
            elif DOCX2TXT_AVAILABLE:
                result = self._process_with_docx2txt(file_path)
            else:
                return WordProcessingResult(
                    success=False,
                    text_content="",
                    tables=[],
                    transactions=[],
                    processing_time=time.time() - start_time,
                    metadata={},
                    error_message="No Word processing library available"
                )
            
            result.processing_time = time.time() - start_time
            return result
            
        except Exception as e:
            processing_time = time.time() - start_time
            error_msg = f"Word document processing failed: {str(e)}"
            self.logger.error(error_msg, exc_info=True)
            
            return WordProcessingResult(
                success=False,
                text_content="",
                tables=[],
                transactions=[],
                processing_time=processing_time,
                metadata={},
                error_message=error_msg
            )
    
    def _process_with_python_docx(self, file_path: str) -> WordProcessingResult:
        """Process Word document using python-docx library"""
        try:
            doc = Document(file_path)
            
            # Extract text content
            text_content = self._extract_text_from_docx(doc)
            
            # Extract tables
            tables = self._extract_tables_from_docx(doc)
            
            # Analyze content for banking transactions
            transactions = self._extract_transactions_from_content(text_content, tables)
            
            # Calculate metadata
            metadata = {
                'method': 'python-docx',
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                'extracted_tables': len(tables),
                'text_length': len(text_content),
                'banking_keywords_found': self._count_banking_keywords(text_content),
                'date_matches': len(self._find_dates_in_text(text_content)),
                'amount_matches': len(self._find_amounts_in_text(text_content)),
                'file_size': os.path.getsize(file_path)
            }
            
            success = len(transactions) > 0 or len(tables) > 0 or len(text_content) > 100
            
            return WordProcessingResult(
                success=success,
                text_content=text_content,
                tables=tables,
                transactions=transactions,
                processing_time=0.0,  # Will be set by caller
                metadata=metadata
            )
            
        except Exception as e:
            raise Exception(f"python-docx processing failed: {e}")
    
    def _process_with_docx2txt(self, file_path: str) -> WordProcessingResult:
        """Process Word document using docx2txt library (fallback)"""
        try:
            # Extract text only (docx2txt doesn't handle tables well)
            text_content = docx2txt.process(file_path)
            
            # Analyze text for potential table-like structures
            tables = self._extract_tables_from_text(text_content)
            
            # Extract transactions from text
            transactions = self._extract_transactions_from_content(text_content, tables)
            
            metadata = {
                'method': 'docx2txt',
                'text_length': len(text_content),
                'extracted_tables': len(tables),
                'banking_keywords_found': self._count_banking_keywords(text_content),
                'date_matches': len(self._find_dates_in_text(text_content)),
                'amount_matches': len(self._find_amounts_in_text(text_content)),
                'file_size': os.path.getsize(file_path)
            }
            
            success = len(transactions) > 0 or len(tables) > 0 or len(text_content) > 100
            
            return WordProcessingResult(
                success=success,
                text_content=text_content,
                tables=tables,
                transactions=transactions,
                processing_time=0.0,  # Will be set by caller
                metadata=metadata
            )
            
        except Exception as e:
            raise Exception(f"docx2txt processing failed: {e}")
    
    def _extract_text_from_docx(self, doc: Document) -> str:
        """Extract all text content from Word document"""
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        return '\n'.join(text_parts)
    
    def _extract_tables_from_docx(self, doc: Document) -> List[WordTable]:
        """Extract tables from Word document using python-docx"""
        tables = []
        
        for table_idx, table in enumerate(doc.tables):
            try:
                # Extract table data
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_data.append(cell_text)
                    table_data.append(row_data)
                
                if not table_data or len(table_data) < self.quality_thresholds['min_table_rows']:
                    continue
                
                # Calculate table metrics
                row_count = len(table_data)
                col_count = len(table_data[0]) if table_data else 0
                
                # Check if table looks like banking data
                confidence = self._calculate_table_confidence(table_data)
                
                if confidence > 0.3:  # Minimum confidence threshold
                    has_header = self._detect_table_header(table_data)
                    
                    word_table = WordTable(
                        data=table_data,
                        row_count=row_count,
                        col_count=col_count,
                        confidence=confidence,
                        table_index=table_idx,
                        has_header=has_header
                    )
                    tables.append(word_table)
                    
                    self.logger.debug(f"Extracted table {table_idx}: {row_count}x{col_count}, confidence: {confidence:.2f}")
                
            except Exception as e:
                self.logger.debug(f"Failed to process table {table_idx}: {e}")
                continue
        
        return tables
    
    def _extract_tables_from_text(self, text: str) -> List[WordTable]:
        """Extract table-like structures from plain text (fallback method)"""
        tables = []
        lines = text.split('\n')
        
        # Look for table-like patterns (multiple columns separated by tabs or multiple spaces)
        potential_tables = []
        current_table = []
        
        for line in lines:
            line = line.strip()
            if not line:
                if current_table and len(current_table) >= self.quality_thresholds['min_table_rows']:
                    potential_tables.append(current_table)
                current_table = []
                continue
            
            # Check if line looks like a table row (multiple columns)
            if '\t' in line:
                columns = [col.strip() for col in line.split('\t')]
            elif '  ' in line:  # Multiple spaces
                columns = [col.strip() for col in re.split(r'\s{2,}', line) if col.strip()]
            else:
                columns = [line]
            
            if len(columns) >= self.quality_thresholds['min_table_cols']:
                current_table.append(columns)
            else:
                if current_table and len(current_table) >= self.quality_thresholds['min_table_rows']:
                    potential_tables.append(current_table)
                current_table = []
        
        # Process potential tables
        for table_idx, table_data in enumerate(potential_tables):
            confidence = self._calculate_table_confidence(table_data)
            
            if confidence > 0.3:
                has_header = self._detect_table_header(table_data)
                
                word_table = WordTable(
                    data=table_data,
                    row_count=len(table_data),
                    col_count=len(table_data[0]) if table_data else 0,
                    confidence=confidence,
                    table_index=table_idx,
                    has_header=has_header
                )
                tables.append(word_table)
        
        return tables
    
    def _calculate_table_confidence(self, table_data: List[List[str]]) -> float:
        """Calculate confidence that table contains banking data"""
        if not table_data:
            return 0.0
        
        confidence = 0.0
        total_cells = sum(len(row) for row in table_data)
        
        # Check for banking-related headers
        if table_data:
            first_row = table_data[0]
            header_score = 0.0
            for cell in first_row:
                cell_lower = cell.lower()
                for pattern in self.table_headers:
                    if re.search(pattern, cell_lower, re.IGNORECASE):
                        header_score += 1
                        break
            
            if len(first_row) > 0:
                confidence += (header_score / len(first_row)) * 0.4
        
        # Check for date patterns
        date_matches = 0
        for row in table_data:
            for cell in row:
                if self._looks_like_date(cell):
                    date_matches += 1
        
        if total_cells > 0:
            confidence += (date_matches / total_cells) * 0.3
        
        # Check for amount patterns
        amount_matches = 0
        for row in table_data:
            for cell in row:
                if self._looks_like_amount(cell):
                    amount_matches += 1
        
        if total_cells > 0:
            confidence += (amount_matches / total_cells) * 0.3
        
        return min(confidence, 1.0)
    
    def _detect_table_header(self, table_data: List[List[str]]) -> bool:
        """Detect if table has a header row"""
        if not table_data or len(table_data) < 2:
            return False
        
        first_row = table_data[0]
        
        # Check if first row contains typical header words
        header_indicators = 0
        for cell in first_row:
            cell_lower = cell.lower()
            for pattern in self.table_headers:
                if re.search(pattern, cell_lower, re.IGNORECASE):
                    header_indicators += 1
                    break
        
        # Header likely if more than half the cells match header patterns
        return header_indicators > len(first_row) / 2
    
    def _extract_transactions_from_content(self, text: str, tables: List[WordTable]) -> List[Dict[str, Any]]:
        """Extract banking transactions from text and tables"""
        transactions = []
        
        # Extract from tables first (more structured)
        for table in tables:
            table_transactions = self._extract_transactions_from_table(table)
            transactions.extend(table_transactions)
        
        # Extract from text if no table transactions found
        if not transactions:
            text_transactions = self._extract_transactions_from_text_content(text)
            transactions.extend(text_transactions)
        
        return transactions
    
    def _extract_transactions_from_table(self, table: WordTable) -> List[Dict[str, Any]]:
        """Extract transactions from a table structure"""
        transactions = []
        
        if not table.data:
            return transactions
        
        # Determine column mapping
        header_row = table.data[0] if table.has_header else None
        data_rows = table.data[1:] if table.has_header else table.data
        
        column_mapping = self._map_table_columns(header_row, table.data)
        
        for row_idx, row in enumerate(data_rows):
            try:
                transaction = {}
                
                # Extract fields based on column mapping
                if column_mapping.get('date') is not None and column_mapping['date'] < len(row):
                    date_value = row[column_mapping['date']].strip()
                    if date_value:
                        transaction['date'] = self._parse_date(date_value)
                
                if column_mapping.get('description') is not None and column_mapping['description'] < len(row):
                    desc_value = row[column_mapping['description']].strip()
                    if desc_value:
                        transaction['description'] = desc_value
                
                if column_mapping.get('amount') is not None and column_mapping['amount'] < len(row):
                    amount_value = row[column_mapping['amount']].strip()
                    if amount_value:
                        transaction['amount'] = self._parse_amount(amount_value)
                
                if column_mapping.get('balance') is not None and column_mapping['balance'] < len(row):
                    balance_value = row[column_mapping['balance']].strip()
                    if balance_value:
                        transaction['balance'] = self._parse_amount(balance_value)
                
                # Only include transactions with at least date or description or amount
                if transaction.get('date') or transaction.get('description') or transaction.get('amount') is not None:
                    transaction['source'] = f'table_{table.table_index}_row_{row_idx}'
                    transaction['confidence'] = self._calculate_transaction_confidence(transaction)
                    transactions.append(transaction)
                
            except Exception as e:
                self.logger.debug(f"Failed to process table row {row_idx}: {e}")
                continue
        
        return transactions
    
    def _extract_transactions_from_text_content(self, text: str) -> List[Dict[str, Any]]:
        """Extract transactions from plain text content"""
        transactions = []
        lines = text.split('\n')
        
        for line_idx, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            
            # Look for lines that contain both dates and amounts
            dates = self._find_dates_in_text(line)
            amounts = self._find_amounts_in_text(line)
            
            if dates and amounts:
                transaction = {
                    'date': self._parse_date(dates[0]),
                    'description': line,
                    'amount': self._parse_amount(amounts[0]),
                    'source': f'text_line_{line_idx}',
                    'confidence': 0.6  # Lower confidence for text extraction
                }
                transactions.append(transaction)
        
        return transactions
    
    def _map_table_columns(self, header_row: Optional[List[str]], table_data: List[List[str]]) -> Dict[str, Optional[int]]:
        """Map table columns to banking data fields"""
        mapping = {
            'date': None,
            'description': None,
            'amount': None,
            'balance': None,
            'reference': None
        }
        
        if not header_row:
            # Try to infer from data patterns
            if table_data and len(table_data) > 0:
                sample_row = table_data[0]
                for col_idx, cell in enumerate(sample_row):
                    if self._looks_like_date(cell) and mapping['date'] is None:
                        mapping['date'] = col_idx
                    elif self._looks_like_amount(cell) and mapping['amount'] is None:
                        mapping['amount'] = col_idx
                    elif len(cell) > 10 and mapping['description'] is None:  # Longer text likely description
                        mapping['description'] = col_idx
            return mapping
        
        # Map based on header text
        for col_idx, header in enumerate(header_row):
            header_lower = header.lower().strip()
            
            # Date column
            if any(re.search(pattern, header_lower, re.IGNORECASE) for pattern in [r'fecha', r'date']):
                mapping['date'] = col_idx
            
            # Description column
            elif any(re.search(pattern, header_lower, re.IGNORECASE) for pattern in [r'descripci[oó]n', r'description', r'concepto', r'concept']):
                mapping['description'] = col_idx
            
            # Amount column
            elif any(re.search(pattern, header_lower, re.IGNORECASE) for pattern in [r'importe', r'amount', r'monto', r'valor']):
                mapping['amount'] = col_idx
            
            # Balance column
            elif any(re.search(pattern, header_lower, re.IGNORECASE) for pattern in [r'saldo', r'balance']):
                mapping['balance'] = col_idx
            
            # Reference column
            elif any(re.search(pattern, header_lower, re.IGNORECASE) for pattern in [r'referencia', r'reference', r'ref']):
                mapping['reference'] = col_idx
        
        return mapping
    
    def _looks_like_date(self, value: str) -> bool:
        """Check if value looks like a date"""
        if not value or len(value.strip()) < 6:
            return False
        
        for pattern in self.banking_patterns['date']:
            if re.search(pattern, value, re.IGNORECASE):
                return True
        
        return False
    
    def _looks_like_amount(self, value: str) -> bool:
        """Check if value looks like a monetary amount"""
        if not value:
            return False
        
        for pattern in self.banking_patterns['amount']:
            if re.search(pattern, value):
                return True
        
        return False
    
    def _find_dates_in_text(self, text: str) -> List[str]:
        """Find all date-like patterns in text"""
        dates = []
        for pattern in self.banking_patterns['date']:
            matches = re.findall(pattern, text, re.IGNORECASE)
            dates.extend(matches)
        return dates
    
    def _find_amounts_in_text(self, text: str) -> List[str]:
        """Find all amount-like patterns in text"""
        amounts = []
        for pattern in self.banking_patterns['amount']:
            matches = re.findall(pattern, text)
            amounts.extend(matches)
        return amounts
    
    def _count_banking_keywords(self, text: str) -> int:
        """Count banking-related keywords in text"""
        count = 0
        text_lower = text.lower()
        
        for pattern in self.banking_patterns['transaction_keywords']:
            matches = re.findall(pattern, text_lower, re.IGNORECASE)
            count += len(matches)
        
        return count
    
    def _parse_date(self, date_str: str) -> Optional[str]:
        """Parse date string to standardized format"""
        if not date_str:
            return None
        
        try:
            # Try common date formats
            for fmt in ['%d/%m/%Y', '%d-%m-%Y', '%Y/%m/%d', '%Y-%m-%d', '%d/%m/%y', '%d-%m-%y']:
                try:
                    parsed = datetime.strptime(date_str.strip(), fmt)
                    return parsed.strftime('%Y-%m-%d')
                except ValueError:
                    continue
            
            # Return original if parsing fails
            return date_str.strip()
            
        except Exception:
            return date_str.strip()
    
    def _parse_amount(self, amount_str: str) -> Optional[float]:
        """Parse amount string to float"""
        if not amount_str:
            return None
        
        try:
            # Clean the amount string
            cleaned = amount_str.replace('$', '').replace('€', '').replace(',', '').strip()
            
            # Handle negative amounts in parentheses
            if cleaned.startswith('(') and cleaned.endswith(')'):
                cleaned = '-' + cleaned[1:-1]
            
            return float(cleaned)
            
        except Exception:
            return None
    
    def _calculate_transaction_confidence(self, transaction: Dict[str, Any]) -> float:
        """Calculate confidence score for a transaction"""
        confidence = 0.0
        
        if transaction.get('date'):
            confidence += 0.3
        
        if transaction.get('description'):
            confidence += 0.3
            if len(transaction['description']) > 10:
                confidence += 0.1
        
        if transaction.get('amount') is not None:
            confidence += 0.3
        
        if transaction.get('balance') is not None:
            confidence += 0.1
        
        return min(confidence, 1.0)
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats"""
        return ['.docx', '.doc']
    
    def validate_file_format(self, file_path: str) -> bool:
        """Validate if file format is supported"""
        file_ext = Path(file_path).suffix.lower()
        return file_ext in ['.docx']  # Only .docx is reliably supported